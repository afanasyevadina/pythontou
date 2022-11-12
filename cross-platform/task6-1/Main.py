#!/usr/bin/env python3
# coding=utf-8

from docx import Document
import xlsxwriter

characters = [
    {
        'name': 'Ричард Хендрикс',
        'about': 'Создатель революционного алгоритма сжатия, основатель компании Pied Piper (встречаются переводы Пегий дудочник и Крысолов).',
        'picture': 'images/richard.jpeg'
    },
    {
        'name': 'Бертрам Гилфойл',
        'about': 'Системный инженер в компании Pied Piper. Нелегальный иммигрант из Канады. Убежденный сатанист. Единственный из компании, у кого есть постоянная девушка, с которой, впрочем, находится в свободных отношениях. Чрезвычайно саркастичен, особенно по отношению к Динешу.',
        'picture': 'images/gilfoyle.jpeg'
    },
    {
        'name': 'Динеш Чугтай',
        'about': 'JAVA-программист в Pied Piper. "Пакистанский Дэнзел Вашингтон", объект постоянный подколов Гилфойла, которому катастрофически не везет с девушками.',
        'picture': 'images/dinesh.jpeg'
    },
    {
        'name': 'Дональд Данн (Джаред)',
        'about': 'Был правой рукой Гевина Белсона пока не увлекся идеей Ричарда Хендрикса и не покинул компанию ради Pied Piper. Невероятно предан Ричарду. Живет в гараже бизнес-инкубатора, поскольку не может выгнать жильца из собственной квартиры.',
        'picture': 'images/jared.jpeg'
    },
    {
        'name': 'Моника Холл',
        'about': 'Самый молодой ассоциированный партнер Raviga Capital, хорошая подруга Ричарда Хендрикса, которая верит в его продукт и всеми силами пытается помочь в его продвижении.',
        'picture': 'images/monica.jpeg'
    },
    {
        'name': 'Лори Брим',
        'about': 'Управляющий партнер компании Raviga Capital.',
        'picture': 'images/laurie.jpeg'
    },
    {
        'name': 'Нельсон Бигетти (Башка)',
        'about': 'Сотрудник Hooli, друг Ричарда Хендрикса. Постоянно попадает в нелепые ситуации из-за своей недалекости и наивности.',
        'picture': 'images/bighetti.jpeg'
    },
    {
        'name': 'Эрлих Бахман',
        'about': 'Владелец бизнес-инкубатора для программистов, в котором работают и живут сотрудники Pied Piper. Жильцы не платят за свое проживание, но взамен отдают 10% своих компаний Эрлиху. Курит много травки.',
        'picture': 'images/erlich.jpeg'
    },
    {
        'name': 'Гевин Белсон',
        'about': 'Генеральный директор компании Hooli. Вплоть до четвертого сезона главный враг Ричарда Хендрикса, пытающийся присвоить программные продукты компании Pied Piper.',
        'picture': 'images/gavin.jpeg'
    },
    {
        'name': 'Денпок Сингх',
        'about': 'Гуру, духовный наставник Гевина Белсона, который сильно влияет на его решения по рабочим вопросам.',
        'picture': 'images/singh.jpeg'
    },
    {
        'name': 'Джек Баркер',
        'about': 'Топ-менеджер. Одно время был гендиректором Pied Piper, но отвратительно справлялся со своей работой и был вытеснен с поста Ричардом Хендриксом и его командой.',
        'picture': 'images/barker.jpeg'
    },
    {
        'name': 'Дзан Янг',
        'about': 'Программист из Китая, живущий в бизнес-инкубаторе Эрлиха Бахмана. Все его проекты провальные, но Эрлих не может выгнать Дзан Янга из-за спеифического законодательства в штате Калифорния. Единственный из обитателей инкубатора, кто не связан с Pied Piper.',
        'picture': 'images/yang.jpeg'
    },
    {
        'name': 'Карла Уолтон',
        'about': 'Инженер, нанятый Pied Piper. Остальные сотрудники навязчиво пытались подружить ее с Моникой. Проработала в компании недолго.',
        'picture': 'images/carla.jpeg'
    },
    {
        'name': 'Винни',
        'about': 'Инженер из Фейсбука, с которой Ричард пытался построить отношения, пока не узнал, что она ставит в программе пробелы вместо Tab\'ов.',
        'picture': 'images/winnie.jpeg'
    },
    {
        'name': 'Питер Грегори',
        'about': 'Эксцентричный миллиардер, инвестор. Был главой компании Raviga Capial — главным инвестором в Pied Piper. После его внезапной трагической смерти совет директоров назначил Лори Брим управляющим партнером компании.',
        'picture': 'images/peter.jpeg'
    },
    {
        'name': 'Расс Ханнеман',
        'about': 'Создатель интернет-радио, эксцентричный член клуба трех запятых, миллиардер. Владелец марки текилы "Tres commas".',
        'picture': 'images/russ.jpeg'
    },
    {
        'name': 'Рон ЛаФламм',
        'about': 'Юрист, периодически помогает советами Ричарду Хендриксу.',
        'picture': 'images/ron.jpeg'
    },
    {
        'name': 'Эван Спиридакис',
        'about': 'Администратор и главный офис-менеджер в компании Raviga.',
        'picture': 'images/evan.jpeg'
    }
]

# word

document = Document()
document.add_heading('Кремниевая долина - персонажи', 0)

for character in characters:
    p = document.add_heading(character['name'], level=1)
    document.add_paragraph('')
    document.add_picture(character['picture'])
    document.add_paragraph(character['about'])
    document.add_page_break()

document.save('Silicon Valley.docx')

# excel

try:
    my_file = 'Silicon Valley.xlsx'  # Имя файла
    book = xlsxwriter.Workbook(my_file)  # Создание файла
    for character in characters:
        sheet = book.add_worksheet(character['name'])  # Добавление в него книги
        sheet.set_column('A:A', 50)  # Установка ширины колонки
        bold = book.add_format({'bold': True})  # Формат жирного текста
        text_format = book.add_format({'text_wrap': True})
        sheet.write('A1', character['name'], bold)  # Выдача текста в ячейку
        sheet.write('A2', character['about'], text_format)  # Выдача жирного текста в ячейку
        sheet.insert_image('B1', character['picture'], {'x_offset': 10})  # Вставка в ячейку картинки
    book.close()  # Закрытие файла
except Exception as a:  # Обработка ошибок
    print("Error!")
    print(a)
